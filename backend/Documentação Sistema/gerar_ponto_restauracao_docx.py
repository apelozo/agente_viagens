# -*- coding: utf-8 -*-
"""Gera Ponto de Restauração v1.1.docx a partir do Markdown homónimo."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches


def set_cell_shading(cell, fill: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def is_table_separator(line: str) -> bool:
    s = line.strip()
    if not s.startswith("|"):
        return False
    return bool(re.match(r"^\|[\s\-:|]+\|\s*$", s))


def parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def add_mixed_to_paragraph(p, line: str) -> None:
    if not line.strip():
        return
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", line)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            p.add_run(part)


def add_mixed_paragraph(doc, line: str, style=None):
    p = doc.add_paragraph(style=style)
    add_mixed_to_paragraph(p, line)
    return p


def flush_table(doc, rows: list[list[str]]) -> None:
    if len(rows) < 1:
        return
    if len(rows) >= 2 and all(re.match(r"^[\s\-:]+$", c or "") for c in rows[1]):
        rows = [rows[0]] + rows[2:]
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell = tbl.rows[ri].cells[ci]
            text = row[ci] if ci < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            add_mixed_to_paragraph(p, text)
            if ri == 0:
                for run in p.runs:
                    run.bold = True
                set_cell_shading(cell, "D9E2F3")


def main():
    base = Path(__file__).resolve().parent
    md_path = base / "Ponto de Restauração v1.1.md"
    out_path = base / "Ponto de Restauração v1.1.docx"

    if not md_path.is_file():
        print(f"Ficheiro não encontrado: {md_path}", file=sys.stderr)
        sys.exit(1)

    lines = md_path.read_text(encoding="utf-8").splitlines()

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    i = 0
    in_code = False
    code_lines: list[str] = []

    def flush_code():
        nonlocal code_lines
        if not code_lines:
            return
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run("\n".join(code_lines))
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        code_lines = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_code()
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if line.strip() == "---":
            doc.add_paragraph()
            i += 1
            continue

        # Bloco de tabela (linhas consecutivas que começam por |)
        stripped = line.strip()
        if stripped.startswith("|") and stripped.count("|") >= 2:
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_line = lines[i]
                if is_table_separator(row_line):
                    i += 1
                    continue
                rows.append(parse_table_row(row_line))
                i += 1
            flush_table(doc, rows)
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue

        m = re.match(r"^[\-\*]\s+(.*)", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_mixed_to_paragraph(p, m.group(1))
            i += 1
            continue

        if line.strip() == "":
            doc.add_paragraph()
            i += 1
            continue

        add_mixed_paragraph(doc, line)
        i += 1

    if in_code and code_lines:
        flush_code()

    # Título principal (mesmo formato tipo “capa” do v1.0): centrado, cor azul escuro
    for p in doc.paragraphs:
        if "Ponto de restauração" in (p.text or ""):
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(20)
                r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            break

    doc.save(out_path)
    print(f"OK: {out_path}")


if __name__ == "__main__":
    main()
